from PIL import Image
from langdetect import detect, DetectorFactory
from langdetect.lang_detect_exception import LangDetectException
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import random
import os
import time
import logging
import pytesseract #This is the package we use to convert image to text
import platform
import subprocess
import shutil

# 配置tesseract路径
def configure_tesseract():
    """
    智能配置tesseract可执行文件路径
    """
    # 常见的tesseract路径
    common_paths = []
    
    # 根据操作系统设置常见路径
    system = platform.system().lower()
    
    if system == 'windows':
        common_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Users\{}\AppData\Local\Tesseract-OCR\tesseract.exe'.format(os.getenv('USERNAME', '')),
        ]
    elif system == 'linux':
        common_paths = [
            '/usr/bin/tesseract',
            '/usr/local/bin/tesseract',
            '/opt/homebrew/bin/tesseract',
        ]
    elif system == 'darwin':  # macOS
        common_paths = [
            '/usr/local/bin/tesseract',
            '/opt/homebrew/bin/tesseract',
            '/usr/bin/tesseract',
        ]
    
    # 方法1: 检查是否已经在PATH中
    tesseract_path = shutil.which('tesseract')
    if tesseract_path:
        logger.info(f"Found tesseract in PATH: {tesseract_path}")
        return tesseract_path
    
    # 方法2: 检查常见路径
    for path in common_paths:
        if os.path.exists(path):
            logger.info(f"Found tesseract at: {path}")
            return path
    
    # 方法3: 尝试使用系统命令查找
    try:
        if system == 'linux' or system == 'darwin':
            result = subprocess.run(['which', 'tesseract'], capture_output=True, text=True)
            if result.returncode == 0:
                path = result.stdout.strip()
                logger.info(f"Found tesseract using 'which': {path}")
                return path
        elif system == 'windows':
            result = subprocess.run(['where', 'tesseract'], capture_output=True, text=True)
            if result.returncode == 0:
                path = result.stdout.strip().split('\n')[0]
                logger.info(f"Found tesseract using 'where': {path}")
                return path
    except Exception as e:
        logger.warning(f"Could not find tesseract using system command: {e}")
    
    # 如果都找不到，返回None
    logger.warning("Could not find tesseract executable. Please install tesseract-ocr package.")
    return None

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 设置tesseract路径
tesseract_path = configure_tesseract()
#This is crucial for the tesseract to work
#if you find the tesseract not working, please check the tesseract_path
#using the following command in the terminal: which tesseract
#after you find the tesseract_path, you can set the tesseract_path
#using the following code:
#pytesseract.pytesseract.tesseract_cmd = r'<full_path_to_your_tesseract_executable>'
#and then you can use the tesseract to work
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
    logger.info(f"Tesseract configured at: {tesseract_path}")
else:
    logger.error("Tesseract not found! Please install tesseract-ocr package.")
    logger.error("Ubuntu/Debian: sudo apt install tesseract-ocr")
    logger.error("macOS: brew install tesseract")
    logger.error("Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki")

# 验证tesseract是否正常工作
try:
    # 测试tesseract是否可用
    available_langs = pytesseract.get_languages(config='')
    logger.info(f"Available tesseract languages: {available_langs}")
except Exception as e:
    logger.error(f"Tesseract verification failed: {e}")

# 设置随机种子
current_time = int(time.time() * 1000000)  
DetectorFactory.seed = current_time
random.seed(current_time)  
logger.info(f"Language detection seed set to: {current_time}")

LANG_MAP = {
    'zh-cn': 'chi_sim',     
    'zh': 'chi_sim',        
    'en': 'eng',            
    'ja': 'jpn',            
    'ko': 'kor',            
    'fr': 'fra',            
    'de': 'deu',            
    'es': 'spa',            
    'it': 'ita',            
    'pt': 'por',            
    'ru': 'rus',           
    'ar': 'ara',           
    'th': 'tha',            
    'vi': 'vie',            
}

# 语言名称映射表
LANG_NAMES = {
    'chi_sim': '简体中文',
    'chi_tra': '繁体中文',
    'eng': '英语',
    'jpn': '日语',
    'kor': '韩语',
    'fra': '法语',
    'deu': '德语',
    'spa': '西班牙语',
    'ita': '意大利语',
    'por': '葡萄牙语',
    'rus': '俄语',
    'ara': '阿拉伯语',
    'tha': '泰语',
    'vie': '越南语',
}

def get_tesseract_lang(lang_code):
    """
    转换语言代码为Tesseract格式
    """
    # 如果已经是Tesseract格式，直接返回
    if lang_code in LANG_NAMES:
        return lang_code
    
    # 如果是langdetect格式，转换
    if lang_code in LANG_MAP:
        return LANG_MAP[lang_code]
    
    # 默认返回英语
    return 'eng'

def extract_text_with_tesseract(image_path, lang='eng'):
    """
    使用Tesseract OCR提取文本
    """
    try:
        start_time = time.time()
        
        # 打开图像
        image = Image.open(image_path)
        
        # 转换语言代码
        tesseract_lang = get_tesseract_lang(lang)
        
        # 执行OCR识别
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(image, lang=tesseract_lang, config=custom_config)
        
        processing_time = time.time() - start_time
        
        logger.info(f"Tesseract processing time: {processing_time:.2f}s, language: {tesseract_lang}")
        return text.strip(), processing_time
        
    except Exception as e:
        logger.error(f"Tesseract extraction failed: {e}")
        return None, 0.0

def detect_language(text):
    """
    检测文本的语言
    :param text: 要检测的文本
    :return: 检测到的语言代码（Tesseract格式）
    """
    if not text or len(text.strip()) < 3:
        return 'eng'  # 默认返回英语

    try:
        # 使用langdetect检测语言
        detected_lang = detect(text)
        print(f"检测到的语言: {detected_lang}")

        # 转换为Tesseract语言代码
        tesseract_lang = LANG_MAP.get(detected_lang, 'eng')
        print(f"转换为Tesseract语言代码: {tesseract_lang}")

        return tesseract_lang
    except LangDetectException:
        print("语言检测失败，使用默认语言: eng")
        return 'eng'

def auto_detect_and_ocr(image_path):
    """
    自动检测图片中的语言并进行OCR识别
    :param image_path: 图片路径
    :return: tuple (识别的文本, 检测到的语言代码)
    """
    try:
        # 首先用英语进行初步OCR识别，获取文本样本用于语言检测
        initial_text, _ = extract_text_with_tesseract(image_path, 'eng')

        if not initial_text or not initial_text.strip():
            # 如果英语识别不出内容，尝试中文
            initial_text, _ = extract_text_with_tesseract(image_path, 'chi_sim')

        if not initial_text or not initial_text.strip():
            logger.warning("No text detected in initial OCR")
            return None, None

        print(f"初步OCR结果: {initial_text[:100]}...")

        # 检测语言
        detected_lang = detect_language(initial_text)

        # 如果检测到的语言与初始使用的不同，重新进行OCR
        if detected_lang not in ['eng', 'chi_sim']:
            print(f"重新使用检测到的语言进行OCR: {detected_lang}")
            final_text, _ = extract_text_with_tesseract(image_path, detected_lang)
        else:
            final_text = initial_text

        return final_text, detected_lang

    except Exception as e:
        print(f"自动语言检测OCR失败: {str(e)}")
        return None, None

def image_to_text(image_path, lang=None):
    """
    将图片转换为文本，支持自动语言检测
    :param image_path: 图片路径
    :param lang: 指定语言代码，如果为None则自动检测
    :return: 识别的文本
    """
    try:
        if lang is None or lang == 'auto':
            # 自动检测语言
            text, detected_lang = auto_detect_and_ocr(image_path)
            print(f"自动检测结果 - 语言: {detected_lang}, 文本长度: {len(text) if text else 0}")
            return text
        else:
            # 使用指定语言
            tesseract_lang = get_tesseract_lang(lang)
            text, processing_time = extract_text_with_tesseract(image_path, tesseract_lang)
            print(f"指定语言OCR结果 - 语言: {tesseract_lang}, 文本长度: {len(text) if text else 0}")
            return text

    except Exception as e:
        print(f"图片文字识别失败: {str(e)}")
        return None

def text_to_word(text, output_path):
    """
    将文本保存为Word文档
    :param text: 要保存的文本
    :param output_path: 输出文件路径
    :return: 是否成功
    """
    try:
        doc = Document()
        doc.add_heading('OCR识别结果', 0)
        doc.add_paragraph(text)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"生成Word文档失败: {str(e)}")
        return False

def text_to_pdf(text, output_path):
    """
    将文本保存为PDF文档
    :param text: 要保存的文本
    :param output_path: 输出文件路径
    :return: 是否成功
    """
    try:
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        
        # 设置字体（支持中文）
        try:
            # 尝试注册中文字体
            font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('DejaVu', font_path))
                c.setFont('DejaVu', 12)
            else:
                c.setFont('Helvetica', 12)
        except:
            c.setFont('Helvetica', 12)

        # 添加标题
        c.setFont('Helvetica-Bold', 16)
        c.drawString(50, height - 50, 'OCR Recognition Result')

        # 添加文本内容
        c.setFont('DejaVu' if os.path.exists("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf") else 'Helvetica', 12)

        # 处理文本换行
        lines = text.split('\n')
        y_position = height - 100
        line_height = 20

        for line in lines:
            if y_position < 50:  # 如果页面空间不够，创建新页面
                c.showPage()
                y_position = height - 50
            
            # 处理长行的换行
            if len(line) > 80:
                words = line.split(' ')
                current_line = ''
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + ' '
                    else:
                        c.drawString(50, y_position, current_line.strip())
                        y_position -= line_height
                        current_line = word + ' '
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
                if current_line:
                    c.drawString(50, y_position, current_line.strip())
                    y_position -= line_height
            else:
                c.drawString(50, y_position, line)
                y_position -= line_height
        
        c.save()
        return True
    except Exception as e:
        print(f"生成PDF文档失败: {str(e)}")
        return False

def get_language_name(lang_code):
    """
    获取语言代码对应的中文名称
    :param lang_code: 语言代码
    :return: 语言的中文名称
    """
    return LANG_NAMES.get(lang_code, '未知语言')
